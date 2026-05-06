"""
AI 智能客服系统 - RAG 文档处理管道

完整的文档处理流程：
1. 分块（FabricChunker）
2. 向量化并存入 DashVector
3. 存入 PostgreSQL（用于 BM25）
4. 返回 doc_id

同时提供删除和查询功能。
"""

import uuid
import json
import os
import io
from typing import List, Dict, Any, Optional
from loguru import logger

from sqlalchemy import text

from app.rag.chunker import FabricChunker
from app.rag.vector_store import VectorStore, get_vector_store
from app.rag.retriever import HybridRetriever, get_hybrid_retriever, HybridSearchResult
from app.rag.reranker import DashScopeReranker, get_reranker
from app.config import settings
from app.utils.database import AsyncSessionLocal


class RAGPipeline:
    """RAG 文档处理管道
    
    提供完整的文档处理流程：
    - 文档入库：分块 -> 向量化 -> 存储
    - 文档删除：删除向量 + 删除分块记录
    - 混合检索：向量 + BM25 + RRF 融合
    
    多租户隔离：
    - 所有操作都携带 tenant_id
    - DashVector 使用 tenant_{tenant_id} Collection
    - PostgreSQL 使用 tenant_id 过滤
    """
    
    def __init__(
        self,
        chunker: Optional[FabricChunker] = None,
        vector_store: Optional[VectorStore] = None,
        retriever: Optional[HybridRetriever] = None,
        reranker: Optional[DashScopeReranker] = None,
    ):
        """
        初始化 RAG Pipeline
        
        Args:
            chunker: 文档分块器
            vector_store: 向量存储
            retriever: 混合检索器
            reranker: 重排序器
        """
        self.chunker = chunker or FabricChunker()
        self.vector_store = vector_store
        self.retriever = retriever
        self.reranker = reranker
    
    async def _ensure_components(self):
        """确保所有组件已初始化"""
        if self.vector_store is None:
            self.vector_store = await get_vector_store()
        
        if self.retriever is None:
            self.retriever = await get_hybrid_retriever()
        
        if self.reranker is None and settings.RERANK_ENABLED:
            self.reranker = await get_reranker()
    
    @staticmethod
    def parse_file(file_path: str = None, file_content: bytes = None, file_name: str = None) -> str:
        """
        解析文档文件，提取文本内容
        
        支持格式：txt, md, docx, pdf
        
        Args:
            file_path: 文件路径（与 file_content 二选一）
            file_content: 文件二进制内容
            file_name: 文件名（用于判断格式，仅在使用 file_content 时需要）
            
        Returns:
            提取的文本内容
        """
        # 确定文件后缀
        name = file_name or file_path or ""
        ext = os.path.splitext(name)[1].lower()
        
        if ext in ('.txt', '.md'):
            if file_content:
                return file_content.decode('utf-8', errors='replace')
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        
        elif ext == '.docx':
            try:
                from docx import Document as DocxDocument
            except ImportError:
                raise ImportError("python-docx 未安装，无法解析 docx 文件，请执行: pip install python-docx")
            if file_content:
                doc = DocxDocument(io.BytesIO(file_content))
            else:
                doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n'.join(paragraphs)
        
        elif ext == '.pdf':
            try:
                import pdfplumber
            except ImportError:
                raise ImportError("pdfplumber 未安装，无法解析 pdf 文件，请执行: pip install pdfplumber")
            source = io.BytesIO(file_content) if file_content else file_path
            pages_text = []
            with pdfplumber.open(source) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
            return '\n\n'.join(pages_text)
        
        else:
            # 未知格式，尝试作为文本读取
            logger.warning(f"Unknown file extension '{ext}', trying as plain text")
            if file_content:
                return file_content.decode('utf-8', errors='replace')
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
    
    async def process_file(
        self,
        file_path: str = None,
        file_content: bytes = None,
        file_name: str = None,
        metadata: Optional[Dict[str, Any]] = None,
        tenant_id: int = 0,
        doc_id: Optional[str] = None,
    ) -> str:
        """
        完整的文件处理流程：解析 -> 分块 -> Embedding -> 存储
        
        Args:
            file_path: 文件路径
            file_content: 文件二进制内容
            file_name: 文件名
            metadata: 文档元数据
            tenant_id: 租户 ID
            doc_id: 文档 ID
            
        Returns:
            doc_id
        """
        # 解析文件内容
        filename = file_name or (os.path.basename(file_path) if file_path else "unknown")
        ext = os.path.splitext(filename)[1].lower()
        logger.info(f"[rag-pipeline] Processing file | filename={filename} type={ext}")
        
        try:
            content = self.parse_file(
                file_path=file_path,
                file_content=file_content,
                file_name=file_name,
            )
        except Exception as e:
            logger.error(f"[rag-pipeline] File processing failed | filename={filename} error={type(e).__name__}: {e}", exc_info=True)
            raise
        
        if not content or not content.strip():
            raise ValueError("文件内容为空，无法处理")
        
        metadata = metadata or {}
        # 从文件名推断标题
        if not metadata.get("title") and (file_name or file_path):
            name = file_name or os.path.basename(file_path)
            metadata["title"] = os.path.splitext(name)[0]
        
        return await self.process_document(
            content=content,
            metadata=metadata,
            tenant_id=tenant_id,
            doc_id=doc_id,
        )
    
    async def process_document(
        self,
        content: str,
        metadata: Dict[str, Any],
        tenant_id: int,
        doc_id: Optional[str] = None,
    ) -> str:
        """
        完整的文档处理流程
        
        Args:
            content: 文档文本内容
            metadata: 文档元数据
                - title: 文档标题
                - doc_type: 文档类型（faq / product / guide / general）
                - category: 分类
                - product_id: 关联商品 ID（可选）
            tenant_id: 租户 ID
            doc_id: 指定文档 ID（可选，默认自动生成）
            
        Returns:
            doc_id: 文档 ID
        """
        await self._ensure_components()
        
        # 生成文档 ID
        if not doc_id:
            doc_id = f"doc_{uuid.uuid4().hex[:16]}"
        
        doc_type = metadata.get("doc_type", "general")
        
        logger.info(f"Processing document {doc_id} (type: {doc_type}, tenant: {tenant_id})")
        
        try:
            # 1. 分块
            chunks = self.chunker.chunk_document(
                content=content,
                doc_type=doc_type,
                metadata={
                    **metadata,
                    "document_id": doc_id,
                    "tenant_id": tenant_id,
                }
            )
            
            if not chunks:
                logger.warning(f"No chunks generated for document {doc_id}")
                return doc_id
            
            chunk_count = len(chunks)
            logger.debug(f"[rag-pipeline] Text chunked | filename={metadata.get('title', doc_id)} chunks={chunk_count}")
            logger.info(f"Document {doc_id} chunked into {chunk_count} chunks")
            
            # 2. 存入 PostgreSQL（用于 BM25 检索）
            await self._save_chunks_to_db(chunks, tenant_id, doc_id)
            
            # 3. 向量化并存入 DashVector
            if self.vector_store and self.vector_store._available:
                success = await self.vector_store.upsert_documents(
                    chunks=chunks,
                    tenant_id=tenant_id,
                    doc_id=doc_id,
                )
                if success:
                    logger.info(f"[rag-pipeline] Vectorization completed | filename={metadata.get('title', doc_id)} chunks={chunk_count}")
                else:
                    logger.warning(f"Failed to upsert vectors for document {doc_id}")
            else:
                logger.warning("VectorStore not available, skipping vector upsert")
            
            return doc_id
            
        except Exception as e:
            logger.error(f"[rag-pipeline] File processing failed | filename={metadata.get('title', doc_id)} error={type(e).__name__}: {e}", exc_info=True)
            raise
    
    async def _save_chunks_to_db(
        self,
        chunks: List[Dict[str, Any]],
        tenant_id: int,
        doc_id: str,
    ):
        """
        保存分块到 PostgreSQL
        
        Args:
            chunks: 分块列表
            tenant_id: 租户 ID
            doc_id: 文档 ID
        """
        try:
            async with AsyncSessionLocal() as session:
                for i, chunk in enumerate(chunks):
                    # 确保 metadata 是 JSON 字符串（PostgreSQL JSONB 字段）
                    metadata_value = chunk.get("metadata", {})
                    if isinstance(metadata_value, dict):
                        metadata_json = json.dumps(metadata_value, ensure_ascii=False)
                    else:
                        metadata_json = str(metadata_value)
                    
                    await session.execute(
                        text("""
                            INSERT INTO rag_chunks (
                                chunk_id, tenant_id, document_id, content, 
                                metadata, chunk_index, created_at
                            ) VALUES (
                                :chunk_id, :tenant_id, :document_id, :content,
                                :metadata::jsonb, :chunk_index, NOW()
                            )
                            ON CONFLICT (chunk_id) DO UPDATE SET
                                content = EXCLUDED.content,
                                metadata = EXCLUDED.metadata,
                                chunk_index = EXCLUDED.chunk_index
                        """),
                        {
                            "chunk_id": chunk["chunk_id"],
                            "tenant_id": tenant_id,
                            "document_id": doc_id,
                            "content": chunk["content"],
                            "metadata": metadata_json,
                            "chunk_index": i,
                        }
                    )
                
                await session.commit()
                logger.info(f"Saved {len(chunks)} chunks to PostgreSQL for doc {doc_id}")
                
        except Exception as e:
            logger.error(f"Failed to save chunks to DB: {e}")
            raise
    
    async def delete_document(
        self,
        doc_id: str,
        tenant_id: int,
    ) -> bool:
        """
        删除文档及其所有分块和向量
        
        Args:
            doc_id: 文档 ID
            tenant_id: 租户 ID
            
        Returns:
            是否成功
        """
        await self._ensure_components()
        
        logger.info(f"Deleting document {doc_id} (tenant: {tenant_id})")
        
        try:
            # 1. 删除 DashVector 中的向量
            if self.vector_store and self.vector_store._available:
                vector_deleted = await self.vector_store.delete_document(
                    tenant_id=tenant_id,
                    doc_id=doc_id,
                )
                if vector_deleted:
                    logger.info(f"Deleted vectors for document {doc_id}")
                else:
                    logger.warning(f"Failed to delete vectors for document {doc_id}")
            
            # 2. 删除 PostgreSQL 中的分块记录
            async with AsyncSessionLocal() as session:
                result = await session.execute(
                    text("""
                        DELETE FROM rag_chunks
                        WHERE tenant_id = :tenant_id AND document_id = :document_id
                    """),
                    {
                        "tenant_id": tenant_id,
                        "document_id": doc_id,
                    }
                )
                await session.commit()
                deleted_count = result.rowcount
                logger.info(f"Deleted {deleted_count} chunks from PostgreSQL for doc {doc_id}")
            
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document {doc_id}: {e}")
            return False
    
    async def query(
        self,
        question: str,
        tenant_id: int,
        top_k: int = 5,
        filters: Optional[Dict[str, Any]] = None,
    ) -> List[HybridSearchResult]:
        """
        混合检索查询
        
        Args:
            question: 查询问题
            tenant_id: 租户 ID
            top_k: 返回结果数
            filters: 过滤条件（如 {"doc_type": "faq"}）
            
        Returns:
            检索结果列表
        """
        await self._ensure_components()
        
        if not self.retriever:
            logger.error("Retriever not available")
            return []
        
        # 防御性类型转换：确保 top_k 为整数
        try:
            top_k = int(top_k)
        except (TypeError, ValueError):
            top_k = 5
        
        try:
            # 如果启用了 reranker，先取更多候选文档
            if settings.RERANK_ENABLED and self.reranker:
                # 取 max(top_k, RETRIEVAL_TOP_K) 的 RRF 融合结果送入 reranker
                candidate_k = max(top_k * 4, settings.RETRIEVAL_TOP_K * 2)
                candidates = await self.retriever.search_with_fallback(
                    query=question,
                    tenant_id=tenant_id,
                    top_k=candidate_k,
                    filters=filters,
                )
                
                if candidates:
                    # 将 HybridSearchResult 转为 dict 供 reranker 使用
                    candidate_dicts = [
                        {
                            "chunk_id": r.chunk_id,
                            "content": r.content,
                            "score": r.score,
                            "vector_score": r.vector_score,
                            "bm25_score": r.bm25_score,
                            "vector_rank": r.vector_rank,
                            "bm25_rank": r.bm25_rank,
                            "metadata": r.metadata,
                        }
                        for r in candidates
                    ]
                    
                    rerank_top_k = settings.RERANK_TOP_K if top_k <= settings.RERANK_TOP_K else top_k
                    reranked = await self.reranker.rerank(
                        query=question,
                        documents=candidate_dicts,
                        top_k=rerank_top_k,
                    )
                    
                    # 转换回 HybridSearchResult
                    results = [
                        HybridSearchResult(
                            chunk_id=doc["chunk_id"],
                            content=doc["content"],
                            score=doc.get("rerank_score", doc["score"]),
                            vector_score=doc.get("vector_score"),
                            bm25_score=doc.get("bm25_score"),
                            vector_rank=doc.get("vector_rank"),
                            bm25_rank=doc.get("bm25_rank"),
                            metadata=doc.get("metadata", {}),
                        )
                        for doc in reranked
                    ]
                    
                    logger.info(
                        f"Query '{question[:50]}...' reranked {len(candidates)} -> {len(results)} results"
                    )
                    return results
            
            # 未启用 reranker 或 reranker 不可用，使用原有流程
            results = await self.retriever.search_with_fallback(
                query=question,
                tenant_id=tenant_id,
                top_k=top_k,
                filters=filters,
            )
            
            logger.info(f"Query '{question[:50]}...' returned {len(results)} results")
            return results
            
        except Exception as e:
            logger.error(f"Query failed: {e}")
            return []
    
    async def get_document_chunks(
        self,
        doc_id: str,
        tenant_id: int,
    ) -> List[Dict[str, Any]]:
        """
        获取文档的所有分块
        
        Args:
            doc_id: 文档 ID
            tenant_id: 租户 ID
            
        Returns:
            分块列表
        """
        try:
            async with AsyncSessionLocal() as session:
                result = await session.execute(
                    text("""
                        SELECT 
                            chunk_id, content, metadata, chunk_index, created_at
                        FROM rag_chunks
                        WHERE tenant_id = :tenant_id AND document_id = :document_id
                        ORDER BY chunk_index
                    """),
                    {
                        "tenant_id": tenant_id,
                        "document_id": doc_id,
                    }
                )
                rows = result.fetchall()
                
                return [
                    {
                        "chunk_id": row.chunk_id,
                        "content": row.content,
                        "metadata": row.metadata,
                        "chunk_index": row.chunk_index,
                        "created_at": row.created_at.isoformat() if row.created_at else None,
                    }
                    for row in rows
                ]
                
        except Exception as e:
            logger.error(f"Failed to get document chunks: {e}")
            return []
    
    async def reindex_document(
        self,
        doc_id: str,
        tenant_id: int,
        new_content: str,
        new_metadata: Optional[Dict[str, Any]] = None,
    ) -> bool:
        """
        重新索引文档（删除旧分块和向量，重新处理）
        
        Args:
            doc_id: 文档 ID
            tenant_id: 租户 ID
            new_content: 新文档内容
            new_metadata: 新文档元数据（可选，默认保留原元数据）
            
        Returns:
            是否成功
        """
        try:
            # 获取原元数据
            if new_metadata is None:
                async with AsyncSessionLocal() as session:
                    result = await session.execute(
                        text("""
                            SELECT metadata FROM rag_chunks
                            WHERE tenant_id = :tenant_id AND document_id = :document_id
                            LIMIT 1
                        """),
                        {
                            "tenant_id": tenant_id,
                            "document_id": doc_id,
                        }
                    )
                    row = result.fetchone()
                    if row and row.metadata:
                        new_metadata = row.metadata
                    else:
                        new_metadata = {}
            
            # 删除旧数据
            await self.delete_document(doc_id, tenant_id)
            
            # 重新处理
            await self.process_document(
                content=new_content,
                metadata=new_metadata,
                tenant_id=tenant_id,
                doc_id=doc_id,
            )
            
            logger.info(f"Reindexed document {doc_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to reindex document {doc_id}: {e}")
            return False
    
    async def get_stats(self, tenant_id: int) -> Dict[str, Any]:
        """
        获取 RAG 统计信息
        
        Args:
            tenant_id: 租户 ID
            
        Returns:
            统计信息字典
        """
        try:
            async with AsyncSessionLocal() as session:
                # 分块总数
                result = await session.execute(
                    text("""
                        SELECT COUNT(*) as total_chunks,
                               COUNT(DISTINCT document_id) as total_documents
                        FROM rag_chunks
                        WHERE tenant_id = :tenant_id
                    """),
                    {"tenant_id": tenant_id}
                )
                row = result.fetchone()
                total_chunks = row.total_chunks if row else 0
                total_documents = row.total_documents if row else 0
                
                # 文档类型分布
                result = await session.execute(
                    text("""
                        SELECT 
                            metadata->>'doc_type' as doc_type,
                            COUNT(*) as count
                        FROM rag_chunks
                        WHERE tenant_id = :tenant_id
                        GROUP BY metadata->>'doc_type'
                    """),
                    {"tenant_id": tenant_id}
                )
                doc_type_dist = {
                    row.doc_type or "unknown": row.count 
                    for row in result
                }
                
                # VectorStore 状态
                vector_status = {}
                if self.vector_store:
                    vector_status = await self.vector_store.health_check()
                
                return {
                    "tenant_id": tenant_id,
                    "total_chunks": total_chunks,
                    "total_documents": total_documents,
                    "doc_type_distribution": doc_type_dist,
                    "vector_store": vector_status,
                }
                
        except Exception as e:
            logger.error(f"Failed to get RAG stats: {e}")
            return {
                "tenant_id": tenant_id,
                "total_chunks": 0,
                "total_documents": 0,
                "doc_type_distribution": {},
                "error": str(e),
            }


# 全局 RAGPipeline 实例
_rag_pipeline: Optional[RAGPipeline] = None


async def get_rag_pipeline() -> RAGPipeline:
    """获取 RAGPipeline 实例（单例模式）"""
    global _rag_pipeline
    if _rag_pipeline is None:
        _rag_pipeline = RAGPipeline()
        await _rag_pipeline._ensure_components()
    return _rag_pipeline


def reset_rag_pipeline():
    """重置 RAGPipeline 实例（用于测试）"""
    global _rag_pipeline
    _rag_pipeline = None


# 便捷函数
async def process_document(
    content: str,
    metadata: Dict[str, Any],
    tenant_id: int,
    doc_id: Optional[str] = None,
) -> str:
    """便捷的文档处理函数"""
    pipeline = await get_rag_pipeline()
    return await pipeline.process_document(content, metadata, tenant_id, doc_id)


async def search_knowledge(
    query: str,
    tenant_id: int,
    top_k: int = 5,
    filters: Optional[Dict[str, Any]] = None,
) -> List[HybridSearchResult]:
    """便捷的知识库检索函数"""
    pipeline = await get_rag_pipeline()
    return await pipeline.query(query, tenant_id, top_k, filters)
